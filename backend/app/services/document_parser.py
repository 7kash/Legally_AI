"""
Document parser service for extracting text from PDF and DOCX files
Includes OCR support for scanned PDFs
"""

from pathlib import Path
from typing import Dict, Any
import pdfplumber
from docx import Document
import logging
import sys
import signal
from contextlib import contextmanager

# Configure logging
logger = logging.getLogger(__name__)


class TimeoutException(Exception):
    pass


@contextmanager
def timeout(seconds):
    """Context manager for timing out operations"""
    def timeout_handler(signum, frame):
        raise TimeoutException(f"Operation timed out after {seconds} seconds")

    # Set the signal handler
    old_handler = signal.signal(signal.SIGALRM, timeout_handler)
    signal.alarm(seconds)

    try:
        yield
    finally:
        # Restore the old signal handler
        signal.alarm(0)
        signal.signal(signal.SIGALRM, old_handler)


# OCR support for scanned PDFs
try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
    logger.info("OCR libraries loaded successfully")
except ImportError as e:
    OCR_AVAILABLE = False
    logger.warning(f"OCR libraries not available: {e}")


def extract_text_with_ocr(file_path: str) -> Dict[str, Any]:
    """
    Extract text from scanned PDF using OCR

    Args:
        file_path: Path to PDF file

    Returns:
        dict with extracted text and metadata
    """
    if not OCR_AVAILABLE:
        raise ValueError(
            "OCR libraries not available. Install with: pip install pdf2image pytesseract pillow"
        )

    logger.info(f"[OCR] Starting OCR extraction for: {file_path}")
    sys.stdout.flush()

    try:
        # Convert PDF to images with timeout (60 seconds)
        logger.info(f"[OCR] Converting PDF to images...")
        sys.stdout.flush()

        try:
            with timeout(60):
                images = convert_from_path(file_path)
            logger.info(f"[OCR] Converted {len(images)} pages to images")
            sys.stdout.flush()
        except TimeoutException as e:
            logger.error(f"[OCR] PDF to image conversion timed out: {e}")
            raise ValueError("PDF to image conversion timed out after 60 seconds")

        pages = []
        total_chars = 0

        for i, image in enumerate(images):
            # Perform OCR on each page with timeout (30 seconds per page)
            logger.info(f"[OCR] Processing page {i+1}/{len(images)}...")
            sys.stdout.flush()

            try:
                with timeout(30):
                    # Try with English only first (most reliable)
                    text = pytesseract.image_to_string(image, lang='eng', timeout=20)
                logger.info(f"[OCR] Page {i+1} processed: {len(text)} characters extracted")
                sys.stdout.flush()
            except TimeoutException as e:
                logger.error(f"[OCR] Timeout on page {i+1}: {e}")
                sys.stdout.flush()
                text = ""
            except Exception as e:
                logger.error(f"[OCR] Failed on page {i+1}: {e}")
                sys.stdout.flush()
                text = ""

            if text.strip():
                pages.append(text)
                total_chars += len(text)

        full_text = "\n\n".join(pages)
        avg_chars_per_page = total_chars / len(images) if images else 0

        # OCR quality varies, estimate based on text extracted
        quality_score = min(0.8, avg_chars_per_page / 2000)  # Max 0.8 for OCR

        return {
            "text": full_text,
            "page_count": len(images),
            "has_text": len(full_text.strip()) > 0,
            "format": "pdf",
            "is_scanned": True,
            "quality_score": quality_score,
            "avg_chars_per_page": avg_chars_per_page
        }

    except Exception as e:
        raise ValueError(f"Failed to perform OCR on PDF: {str(e)}")


def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text from PDF file using pdfplumber
    Automatically uses OCR if document appears to be scanned

    Returns:
        dict with:
        - text: extracted text
        - page_count: number of pages
        - has_text: whether text was extracted
        - is_scanned: whether document appears to be scanned
        - quality_score: estimate of scan quality (0-1)
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = []
            total_chars = 0

            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
                    total_chars += len(text)

            full_text = "\n\n".join(pages)

            # Estimate quality
            avg_chars_per_page = total_chars / len(pdf.pages) if pdf.pages else 0

            # Typical printed page has 2000-4000 chars
            # Less than 500 suggests poor OCR or scanned without OCR
            quality_score = min(1.0, avg_chars_per_page / 2000)
            is_scanned = avg_chars_per_page < 500

            # If document appears scanned (little/no text), try OCR
            if is_scanned and OCR_AVAILABLE and avg_chars_per_page < 200:
                logger.warning(f"PDF appears scanned (avg {avg_chars_per_page:.0f} chars/page), attempting OCR...")
                sys.stdout.flush()
                logger.info("Calling extract_text_with_ocr function...")
                sys.stdout.flush()
                result = extract_text_with_ocr(file_path)
                logger.info("extract_text_with_ocr completed successfully")
                sys.stdout.flush()
                return result

            return {
                "text": full_text,
                "page_count": len(pdf.pages),
                "has_text": len(full_text.strip()) > 0,
                "format": "pdf",
                "is_scanned": is_scanned,
                "quality_score": quality_score,
                "avg_chars_per_page": avg_chars_per_page
            }

    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text from DOCX file

    Returns:
        dict with:
        - text: extracted text
        - paragraph_count: number of paragraphs
        - has_text: whether text was extracted
        - quality_score: always 1.0 for DOCX (digital format)
    """
    try:
        doc = Document(file_path)

        paragraphs = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = "\n\n".join(paragraphs)

        return {
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "has_text": len(full_text.strip()) > 0,
            "format": "docx",
            "is_scanned": False,
            "quality_score": 1.0  # DOCX is digital, high quality
        }

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text(file_path: str) -> Dict[str, Any]:
    """
    Extract text from PDF or DOCX file
    Auto-detects format based on extension
    Automatically uses OCR for scanned PDFs

    Args:
        file_path: Path to file

    Returns:
        dict with extracted text and metadata
    """
    # Check if file exists
    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"File not found: {file_path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    file_path_lower = file_path.lower()

    if file_path_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_path)

    elif file_path_lower.endswith('.docx'):
        return extract_text_from_docx(file_path)

    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
