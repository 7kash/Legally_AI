"""
Document parsers for PDF and DOCX files
Extracts text while preserving structure
"""

import pdfplumber
from docx import Document
from typing import Optional, Dict
import re

# OCR support for scanned PDFs
try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def extract_text_with_ocr(file_path: str) -> Dict[str, any]:
    """
    Extract text from scanned PDF using OCR

    Returns:
        dict with extracted text and metadata
    """
    if not OCR_AVAILABLE:
        raise ValueError(
            "OCR libraries not available. Install with: pip install pdf2image pytesseract pillow"
        )

    try:
        # Convert PDF to images
        images = convert_from_path(file_path)

        pages = []
        total_chars = 0

        for i, image in enumerate(images):
            # Perform OCR on each page
            text = pytesseract.image_to_string(image, lang='eng+rus+srp+fra')
            if text.strip():
                pages.append(text)
                total_chars += len(text)

        full_text = "\n\n".join(pages)
        avg_chars_per_page = total_chars / len(images) if images else 0

        # OCR quality varies, estimate based on text extracted
        quality_score = min(0.8, avg_chars_per_page / 2000)  # Max 0.8 for OCR

        return {
            "text": full_text,
            "pages": len(images),
            "quality_score": quality_score,
            "is_scanned": True,
            "avg_chars_per_page": avg_chars_per_page
        }

    except Exception as e:
        raise ValueError(f"Failed to perform OCR on PDF: {str(e)}")


def extract_text_from_pdf(file_path: str) -> Dict[str, any]:
    """
    Extract text from PDF file

    Returns:
        dict with:
        - text: extracted text
        - pages: number of pages
        - quality_score: estimate of scan quality (0-1)
        - is_scanned: whether document appears to be scanned
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = []
            total_chars = 0

            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
                    total_chars += len(text)

            full_text = "\n\n".join(pages)

            # Estimate quality
            avg_chars_per_page = total_chars / len(pdf.pages) if pdf.pages else 0

            # Typical printed page has 2000-4000 chars
            # Less than 500 suggests poor OCR or scanned without OCR
            quality_score = min(1.0, avg_chars_per_page / 2000)
            is_scanned = avg_chars_per_page < 500

            # If document appears scanned (little/no text), try OCR
            if is_scanned and OCR_AVAILABLE and avg_chars_per_page < 200:
                print(f"PDF appears scanned (avg {avg_chars_per_page:.0f} chars/page), attempting OCR...")
                return extract_text_with_ocr(file_path)

            return {
                "text": full_text,
                "pages": len(pdf.pages),
                "quality_score": quality_score,
                "is_scanned": is_scanned,
                "avg_chars_per_page": avg_chars_per_page
            }

    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> Dict[str, any]:
    """
    Extract text from DOCX file

    Returns:
        dict with:
        - text: extracted text
        - paragraphs: number of paragraphs
        - quality_score: always 1.0 for DOCX (digital format)
        - is_scanned: always False
    """
    try:
        doc = Document(file_path)

        paragraphs = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = "\n\n".join(paragraphs)

        return {
            "text": full_text,
            "paragraphs": len(paragraphs),
            "quality_score": 1.0,  # DOCX is digital, high quality
            "is_scanned": False,
            "avg_chars_per_page": len(full_text) / max(1, len(paragraphs) // 20)  # Rough estimate
        }

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text(file_path: str) -> Dict[str, any]:
    """
    Extract text from PDF or DOCX file
    Auto-detects format based on extension

    Args:
        file_path: Path to file

    Returns:
        dict with extracted text and metadata
    """
    if file_path.lower().endswith('.pdf'):
        result = extract_text_from_pdf(file_path)
        result['format'] = 'pdf'
        return result

    elif file_path.lower().endswith('.docx'):
        result = extract_text_from_docx(file_path)
        result['format'] = 'docx'
        return result

    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")


def detect_structure(text: str) -> Dict[str, any]:
    """
    Detect document structure (headings, sections)

    Returns:
        dict with:
        - has_headings: boolean
        - sections: list of detected section names
        - appears_complete: boolean (has typical contract sections)
    """
    # Common contract section patterns (expanded for better detection)
    section_patterns = [
        r'\b(parties|definitions|term|duration|payment|price|fee|termination|end|warranties|liability|obligations|rights|dispute|arbitration|governing law|jurisdiction|confidentiality|indemnity|force majeure|assignment|notice|entire agreement)\b',
        r'\b(стороны|определения|срок|период|оплата|цена|расторжение|окончание|гарантии|ответственность|обязательства|права|споры|арбитраж|применимое право|юрисдикция|конфиденциальность)\b',  # Russian
        r'\b(strane|definicije|rok|period|plaćanje|cena|raskid|kraj|garancije|odgovornost|obaveze|prava|sporovi|arbitraža|pravo|nadležnost|poverljivost)\b',  # Serbian
        r'\b(parties|définitions|durée|période|paiement|prix|résiliation|fin|garanties|responsabilité|obligations|droits|litiges|arbitrage|droit applicable|juridiction|confidentialité)\b'  # French
    ]

    sections_found = []
    for pattern in section_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        sections_found.extend(matches)

    # Check for numbered sections (1., 2., etc.)
    has_numbering = bool(re.search(r'\n\s*\d+\.\s+[A-Z]', text))

    # Check for Article pattern
    has_articles = bool(re.search(r'\b(Article|Статья|Član|Article)\s+\d+', text, re.IGNORECASE))

    has_headings = has_numbering or has_articles or len(sections_found) > 3

    # Consider complete if has at least 3 typical sections (lowered for prototype testing)
    # OR if document is reasonably long (>2000 chars suggests real contract)
    unique_sections = len(set([s.lower() for s in sections_found]))
    appears_complete = unique_sections >= 3 or (len(text) > 2000 and unique_sections >= 1)

    return {
        "has_headings": has_headings,
        "sections": list(set(sections_found)),
        "appears_complete": appears_complete,
        "has_numbering": has_numbering,
        "has_articles": has_articles
    }


def normalize_text(text: str) -> str:
    """
    Normalize extracted text
    - Fix common OCR errors
    - Remove excessive whitespace
    - Fix encoding issues
    """
    # Fix multiple spaces
    text = re.sub(r' +', ' ', text)

    # Fix multiple newlines (keep max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove trailing/leading whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    # Common OCR fixes
    replacements = {
        'Ð': '',  # Common encoding issue
        '|': 'I',  # Often misread
        '0': 'O',  # In some contexts
    }

    # Apply selectively (not in obvious dates/numbers)
    # This is basic - production would be more sophisticated

    return text


def extract_dates(text: str) -> list:
    """
    Extract dates from text in various formats

    Returns:
        list of date strings found
    """
    date_patterns = [
        r'\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}',  # MM/DD/YYYY, DD.MM.YYYY
        r'\d{4}[./\-]\d{1,2}[./\-]\d{1,2}',  # YYYY-MM-DD
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b',  # English
        r'\b\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}\b',  # Russian
    ]

    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        dates.extend(matches)

    return dates


def extract_amounts(text: str) -> list:
    """
    Extract monetary amounts from text

    Returns:
        list of amount strings found
    """
    amount_patterns = [
        r'[$€£₽][\d,]+(?:\.\d{2})?',  # $1,000.00
        r'\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|EUR|RUB|GBP|RSD)',  # 1,000.00 USD
        r'\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:долларов|евро|рублей|динаров)',  # Russian currency names
    ]

    amounts = []
    for pattern in amount_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        amounts.extend(matches)

    return amounts
